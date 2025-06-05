import tempfile
import os
from typing import Dict, Any, Optional, Union
from pathlib import Path
import mimetypes
from dataclasses import dataclass
from loguru import logger

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from bs4 import BeautifulSoup
import json
import xml.etree.ElementTree as ET



@dataclass
class ParsedDocument:
    """Represents a parsed document with metadata"""
    content: str
    metadata: Dict[str, Any]
    source_type: str
    source_path: str
    encoding: Optional[str] = None
    page_count: Optional[int] = None
    error: Optional[str] = None


class DocumentParser:
    """Multi-format document parser with specialized handling"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,  # Will try docx parser
            '.txt': self._parse_text,
            '.md': self._parse_markdown,
            '.csv': self._parse_csv,
            '.xlsx': self._parse_excel,
            '.xls': self._parse_excel,
            '.json': self._parse_json,
            '.xml': self._parse_xml,
            '.html': self._parse_html,
            '.htm': self._parse_html,
            '.py': self._parse_code,
            '.js': self._parse_code,
            '.ts': self._parse_code,
            '.java': self._parse_code,
            '.cpp': self._parse_code,
            '.c': self._parse_code,
            '.h': self._parse_code,
        }
    
    async def parse_file(self, file_path: Union[str, Path]) -> ParsedDocument:
        """Parse a file based on its extension"""
        
        file_path = Path(file_path)
        
        if not file_path.exists():
            return ParsedDocument(
                content="",
                metadata={},
                source_type="error",
                source_path=str(file_path),
                error=f"File not found: {file_path}"
            )
        
        # Detect file type
        file_extension = file_path.suffix.lower()
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        # Get parser function
        parser_func = self.supported_formats.get(file_extension)
        
        if not parser_func:
            # Try to parse as text if no specific parser
            logger.warning(f"No specific parser for {file_extension}, trying text parser")
            parser_func = self._parse_text
        
        try:
            logger.info(f"Parsing {file_path} as {file_extension}")
            result = await parser_func(file_path)
            
            # Add common metadata
            result.metadata.update({
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_extension': file_extension,
                'mime_type': mime_type,
                'parsed_at': pd.Timestamp.now().isoformat()
            })
            
            return result
            
        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {e}")
            return ParsedDocument(
                content="",
                metadata={'file_name': file_path.name},
                source_type="error",
                source_path=str(file_path),
                error=str(e)
            )
    
    async def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """Parse PDF files"""
        
        try:
            content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        content += f"\n--- Page {page_num + 1} ---\n"
                        content += page_text
                
                metadata = {
                    'page_count': page_count,
                    'has_annotations': any(page.annotations for page in pdf_reader.pages if page.annotations),
                }
                
                # Try to extract PDF metadata
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                    })
            
            return ParsedDocument(
                content=content.strip(),
                metadata=metadata,
                source_type="pdf",
                source_path=str(file_path),
                page_count=page_count
            )
            
        except Exception as e:
            raise Exception(f"PDF parsing failed: {e}")
    
    async def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """Parse Word documents"""
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract text from tables
            table_content = ""
            for table in doc.tables:
                table_content += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_content += row_text + "\n"
            
            if table_content:
                content += "\n" + table_content
            
            # Extract metadata
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
            }
            
            # Try to extract document properties
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata.update({
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': props.created.isoformat() if props.created else '',
                    'modified': props.modified.isoformat() if props.modified else '',
                })
            
            return ParsedDocument(
                content=content.strip(),
                metadata=metadata,
                source_type="docx",
                source_path=str(file_path)
            )
            
        except Exception as e:
            raise Exception(f"DOCX parsing failed: {e}")
    
    async def _parse_text(self, file_path: Path) -> ParsedDocument:
        """Parse plain text files"""
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            content = ""
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                raise Exception("Could not decode file with any supported encoding")
            
            return ParsedDocument(
                content=content,
                metadata={'encoding': used_encoding, 'line_count': len(content.split('\n'))},
                source_type="text",
                source_path=str(file_path),
                encoding=used_encoding
            )
            
        except Exception as e:
            raise Exception(f"Text parsing failed: {e}")
    
    async def _parse_markdown(self, file_path: Path) -> ParsedDocument:
        """Parse Markdown files"""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Count headers for metadata
            import re
            headers = re.findall(r'^#+\s+(.+)$', content, re.MULTILINE)
            
            return ParsedDocument(
                content=content,
                metadata={
                    'header_count': len(headers),
                    'headers': headers[:10],  # First 10 headers
                    'line_count': len(content.split('\n'))
                },
                source_type="markdown",
                source_path=str(file_path)
            )
            
        except Exception as e:
            raise Exception(f"Markdown parsing failed: {e}")
    
    async def _parse_csv(self, file_path: Path) -> ParsedDocument:
        """Parse CSV files"""
        
        try:
            # Read CSV with pandas for better handling
            df = pd.read_csv(file_path)
            
            # Convert to string representation
            content = df.to_string(index=False)
            
            # Also create a more readable format
            readable_content = f"CSV Data Summary:\n"
            readable_content += f"Columns: {', '.join(df.columns.tolist())}\n"
            readable_content += f"Shape: {df.shape[0]} rows, {df.shape[1]} columns\n\n"
            readable_content += "Sample Data:\n"
            readable_content += df.head(10).to_string(index=False)
            
            if len(df) > 10:
                readable_content += f"\n\n... and {len(df) - 10} more rows"
            
            return ParsedDocument(
                content=readable_content,
                metadata={
                    'row_count': len(df),
                    'column_count': len(df.columns),
                    'columns': df.columns.tolist(),
                    'dtypes': df.dtypes.to_dict(),
                    'memory_usage': df.memory_usage(deep=True).sum()
                },
                source_type="csv",
                source_path=str(file_path)
            )
            
        except Exception as e:
            raise Exception(f"CSV parsing failed: {e}")
    
    async def _parse_excel(self, file_path: Path) -> ParsedDocument:
        """Parse Excel files"""
        
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            content = "Excel File Content:\n\n"
            sheet_info = {}
            
            for sheet_name, df in excel_data.items():
                content += f"Sheet: {sheet_name}\n"
                content += f"Shape: {df.shape[0]} rows, {df.shape[1]} columns\n"
                content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
                content += df.head(5).to_string(index=False)
                content += "\n\n" + "="*50 + "\n\n"
                
                sheet_info[sheet_name] = {
                    'row_count': len(df),
                    'column_count': len(df.columns),
                    'columns': df.columns.tolist()
                }
            
            return ParsedDocument(
                content=content,
                metadata={
                    'sheet_count': len(excel_data),
                    'sheet_names': list(excel_data.keys()),
                    'sheets': sheet_info
                },
                source_type="excel",
                source_path=str(file_path)
            )
            
        except Exception as e:
            raise Exception(f"Excel parsing failed: {e}")
    
    async def _parse_json(self, file_path: Path) -> ParsedDocument:
        """Parse JSON files"""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Pretty print JSON
            content = json.dumps(data, indent=2, ensure_ascii=False)
            
            # Analyze structure
            def analyze_json(obj, path=""):
                info = {"type": type(obj).__name__}
                
                if isinstance(obj, dict):
                    info["keys"] = list(obj.keys())
                    info["size"] = len(obj)
                elif isinstance(obj, list):
                    info["size"] = len(obj)
                    if obj:
                        info["item_type"] = type(obj[0]).__name__
                
                return info
            
            structure_info = analyze_json(data)
            
            return ParsedDocument(
                content=content,
                metadata={
                    'json_structure': structure_info,
                    'size_bytes': len(content.encode('utf-8'))
                },
                source_type="json",
                source_path=str(file_path)
            )
            
        except Exception as e:
            raise Exception(f"JSON parsing failed: {e}")
    
    async def _parse_xml(self, file_path: Path) -> ParsedDocument:
        """Parse XML files"""
        
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Convert to readable text
            content = f"XML Document\nRoot element: {root.tag}\n\n"
            
            def xml_to_text(element, indent=0):
                text = "  " * indent + f"<{element.tag}>"
                if element.text and element.text.strip():
                    text += f" {element.text.strip()}"
                text += "\n"
                
                for child in element:
                    text += xml_to_text(child, indent + 1)
                
                return text
            
            content += xml_to_text(root)
            
            return ParsedDocument(
                content=content,
                metadata={
                    'root_tag': root.tag,
                    'element_count': len(list(root.iter())),
                    'namespace': root.tag.split('}')[0].strip('{') if '}' in root.tag else None
                },
                source_type="xml",
                source_path=str(file_path)
            )
            
        except Exception as e:
            raise Exception(f"XML parsing failed: {e}")
    
    async def _parse_html(self, file_path: Path) -> ParsedDocument:
        """Parse HTML files"""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text content
            content = soup.get_text(separator='\n', strip=True)
            
            # Extract metadata
            title = soup.find('title')
            meta_description = soup.find('meta', attrs={'name': 'description'})
            
            return ParsedDocument(
                content=content,
                metadata={
                    'title': title.get_text() if title else '',
                    'description': meta_description.get('content', '') if meta_description else '',
                    'link_count': len(soup.find_all('a')),
                    'image_count': len(soup.find_all('img')),
                    'has_tables': len(soup.find_all('table')) > 0
                },
                source_type="html",
                source_path=str(file_path)
            )
            
        except Exception as e:
            raise Exception(f"HTML parsing failed: {e}")
    
    async def _parse_code(self, file_path: Path) -> ParsedDocument:
        """Parse code files"""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Detect language from extension
            language_map = {
                '.py': 'python',
                '.js': 'javascript',
                '.ts': 'typescript',
                '.java': 'java',
                '.cpp': 'cpp',
                '.c': 'c',
                '.h': 'c_header'
            }
            
            language = language_map.get(file_path.suffix.lower(), 'unknown')
            
            # Simple analysis
            lines = content.split('\n')
            non_empty_lines = [line for line in lines if line.strip()]
            comment_lines = [line for line in lines if line.strip().startswith(('#', '//', '/*'))]
            
            return ParsedDocument(
                content=content,
                metadata={
                    'language': language,
                    'line_count': len(lines),
                    'non_empty_lines': len(non_empty_lines),
                    'comment_lines': len(comment_lines),
                    'size_bytes': len(content.encode('utf-8'))
                },
                source_type="code",
                source_path=str(file_path)
            )
            
        except Exception as e:
            raise Exception(f"Code parsing failed: {e}")
    
    async def parse_url_content(self, content: Union[str, bytes], url: str, content_type: str = "html") -> ParsedDocument:
        """Parse content from URL (handles both text and binary content)"""
        
        try:
            # Handle binary content (PDFs, Word docs, etc.)
            if isinstance(content, bytes):
                return await self._parse_binary_url_content(content, url, content_type)
            
            # Handle text content
            if content_type.startswith('text/html'):
                soup = BeautifulSoup(content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                text_content = soup.get_text(separator='\n', strip=True)
                
                # Extract metadata
                title = soup.find('title')
                meta_description = soup.find('meta', attrs={'name': 'description'})
                
                return ParsedDocument(
                    content=text_content,
                    metadata={
                        'url': url,
                        'title': title.get_text() if title else '',
                        'description': meta_description.get('content', '') if meta_description else '',
                        'content_type': content_type
                    },
                    source_type="web",
                    source_path=url
                )
            
            elif content_type.startswith('application/json'):
                data = json.loads(content)
                formatted_content = json.dumps(data, indent=2, ensure_ascii=False)
                
                return ParsedDocument(
                    content=formatted_content,
                    metadata={
                        'url': url,
                        'content_type': content_type,
                        'json_keys': list(data.keys()) if isinstance(data, dict) else []
                    },
                    source_type="web_json",
                    source_path=url
                )
            
            else:
                # Plain text or unknown type
                return ParsedDocument(
                    content=content,
                    metadata={
                        'url': url,
                        'content_type': content_type
                    },
                    source_type="web_text",
                    source_path=url
                )
                
        except Exception as e:
            raise Exception(f"URL content parsing failed: {e}")
    
    async def _parse_binary_url_content(self, content: bytes, url: str, content_type: str) -> ParsedDocument:
        """Parse binary content downloaded from URL"""
        
        try:
            # Determine file extension from URL or content type
            file_extension = None
            
            # Try to get extension from URL
            if '.' in url:
                file_extension = url.split('.')[-1].lower()
            
            # Fallback to content type mapping
            if not file_extension:
                content_type_map = {
                    'application/pdf': 'pdf',
                    'application/msword': 'doc',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                    'application/vnd.ms-excel': 'xls',
                    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
                }
                file_extension = content_type_map.get(content_type.lower(), 'bin')
            
            # Create temporary file with appropriate extension
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}') as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                # Parse the temporary file using existing file parser
                parsed_doc = await self.parse_file(Path(temp_file_path))
                
                if parsed_doc.error:
                    raise Exception(f"Failed to parse binary content: {parsed_doc.error}")
                
                # Update metadata with URL information
                parsed_doc.metadata.update({
                    'url': url,
                    'original_content_type': content_type,
                    'downloaded_size': len(content),
                    'file_extension': file_extension
                })
                
                # Update source information
                parsed_doc.source_type = f"web_{file_extension}"
                parsed_doc.source_path = url
                
                logger.info(f"Successfully parsed {file_extension.upper()} from URL: {len(parsed_doc.content)} characters extracted")
                
                return parsed_doc
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except Exception as cleanup_error:
                    logger.warning(f"Failed to cleanup temporary file {temp_file_path}: {cleanup_error}")
                    
        except Exception as e:
            raise Exception(f"Binary URL content parsing failed: {e}")